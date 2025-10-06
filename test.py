# generate_full_article.py
#ابتدا کتابخانه مورد نیاز را با دستور pip install python-docx نصب کنید

import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def setup_styles(document):
    """
    برای تعریف استایل‌های سفارشی مشابه تمپلیت IEEE
    """
    # استایل برای متن اصلی
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)
    
    # استایل برای عنوان مقاله
    style = document.styles.add_style('ArticleTitle', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(22)
    style.font.bold = True
    p_fmt = style.paragraph_format
    p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fmt.space_after = Pt(12)

    # استایل برای نام نویسندگان
    style = document.styles.add_style('Authors', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    p_fmt = style.paragraph_format
    p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # استایل برای وابستگی سازمانی (Affiliation)
    style = document.styles.add_style('Affiliation', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.font.italic = True
    p_fmt = style.paragraph_format
    p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fmt.space_after = Pt(18)

    # استایل برای تیترهای اصلی (I, II, ...)
    style = document.styles.add_style('Heading1IEEE', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.font.all_caps = True
    p_fmt = style.paragraph_format
    p_fmt.space_before = Pt(12)
    p_fmt.space_after = Pt(4)
    
    # استایل برای تیترهای فرعی (A, B, ...)
    style = document.styles.add_style('Heading2IEEE', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)
    style.font.italic = True
    p_fmt = style.paragraph_format
    p_fmt.space_before = Pt(6)
    p_fmt.space_after = Pt(3)
    
    # استایل برای کپشن جدول
    style = document.styles.add_style('TableCaption', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(8)
    style.font.all_caps = True
    p_fmt = style.paragraph_format
    p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fmt.space_before = Pt(10)
    p_fmt.space_after = Pt(2)

    # استایل برای رفرنس ها
    style = document.styles.add_style('Reference', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Times New Roman'
    style.font.size = Pt(8)
    p_fmt = style.paragraph_format
    p_fmt.left_indent = Inches(0.25)
    p_fmt.first_line_indent = Inches(-0.25)

def add_abstract_section(document, abstract_text):
    p = document.add_paragraph()
    p.add_run('Abstract—').bold = True
    p.add_run(abstract_text)
    p.paragraph_format.space_after = Pt(6)
    
def add_index_terms_section(document, terms):
    p = document.add_paragraph()
    p.add_run('Index Terms—').bold = True
    p.add_run(terms)
    p.paragraph_format.space_after = Pt(18)

def add_equation(paragraph, text, number):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.italic = True
    # Add a tab to push the number to the right
    paragraph.add_run().add_tab()
    paragraph.add_run(f'({number})')
    # Set alignment for the whole paragraph, including the number
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_ieee_table(document, caption_text, headers, data):
    p = document.add_paragraph(caption_text, style='TableCaption')
    
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.add_run(header).bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_data in data:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data
    
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = 1 # Center
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(8)
    
    document.add_paragraph().paragraph_format.space_before = Pt(12)

def create_full_document():
    document = docx.Document()
    setup_styles(document)

    # --- Section: Title, Authors, Affiliation ---
    document.add_paragraph(
        "A Cognitive-Computational Framework for Lifelong Learning: Scaffolding Meta-Learning for Emergent Continual Intelligence (SMECI)", 
        style='ArticleTitle'
    )
    document.add_paragraph("Mani Agah and Saman Ghahreman", style='Authors')
    document.add_paragraph(
        "Department of Engineering, Islamic Azad University, Central Tehran Branch, Tehran, Iran\nmani.agah@iauctb.ac.ir, saman.ghahreman@iauctb.ac.ir", 
        style='Affiliation'
    )

    # --- Section: Abstract and Index Terms ---
    add_abstract_section(document,
        "This paper confronts a foundational challenge in modern artificial intelligence: the stability-plasticity dilemma and "
        "the consequent \"catastrophic forgetting\" exhibited by large-scale models when processing sequential, "
        "non-stationary data streams. The inability of contemporary models to learn incrementally hinders their "
        "deployment in real-world, dynamic applications. In response, we propose a novel framework, Scaffolding "
        "Meta-Learning for Emergent Continual Intelligence (SMECI). Its central innovation lies in the computational "
        "operationalization of \"Cognitive Scaffolding,\" a concept derived from developmental psychology. This framework "
        "utilizes a meta-learning controller (the Scaffolding Network) to guide a dynamically architected primary model "
        "(the Learner Network) through its algorithmically-defined \"Zone of Proximal Development\" (ZPD). By formulating "
        "the scaffolding policy as a problem in meta-reinforcement learning, SMECI learns to dynamically select and apply "
        "a sequence of targeted computational supports—including the introduction of auxiliary objectives, dynamic "
        "attention modulation, and guided structural expansion—to optimize the entire learning trajectory. This proactive, "
        "guidance-oriented methodology, which actively manages learning rather than passively preserving knowledge, "
        "stands in stark contrast to prevailing reactive paradigms that primarily penalize parameter modifications. "
        "It is posited that this framework will yield a more parameter-efficient, adaptive, and robust lifelong learning "
        "agent, demonstrating superior performance on complex continual learning benchmarks by not only mitigating forgetting "
        "but also enhancing forward knowledge transfer."
    )
    add_index_terms_section(document, 
        "Catastrophic Forgetting, Cognitive Scaffolding, Continual Learning, Lifelong Learning, Meta-Learning, Stability-Plasticity Dilemma, Zone of Proximal Development."
    )

    # --- Section I: Introduction ---
    document.add_paragraph('I. INTRODUCTION', style='Heading1IEEE')
    document.add_paragraph('A. The Static Nature of Foundation Models and the Imperative for Lifelong Learning', style='Heading2IEEE')
    document.add_paragraph(
        "The contemporary AI paradigm is predominantly characterized by the training of massive, static models, such as "
        "Large Language Models (LLMs) and vision transformers, on vast, fixed datasets [26, 27, 28]. This "
        "\"train-once, deploy-forever\" methodology, while successful in creating powerful generalist models, engenders a "
        "critical vulnerability known as \"catastrophic forgetting\" or \"catastrophic interference,\" wherein a model's "
        "performance on previously acquired tasks degrades precipitously upon learning a new one. This brittleness arises "
        "because the optimization process overwrites synaptic weights crucial for past knowledge to accommodate new "
        "information. This characteristic fundamentally constrains their utility in dynamic, real-world "
        "environments—such as autonomous robotics, personalized medicine, or financial market analysis—where new knowledge "
        "must be continuously and efficiently assimilated [26, 28]. Consequently, Continual Lifelong Learning (CL) "
        "constitutes not merely a desirable capability but an essential evolutionary advancement for artificial intelligence, "
        "marking the transition from static tools to truly adaptive and intelligent systems [29]."
    )
    document.add_paragraph('B. Reframing the Stability-Plasticity Dilemma', style='Heading2IEEE')
    document.add_paragraph(
        "The central theoretical impediment in CL is the stability-plasticity dilemma [16]: a model must possess sufficient "
        "stability to retain past knowledge while remaining adequately plastic to acquire new information. Extant CL "
        "methodologies can be conceptualized as attempts to algorithmically manage this trade-off. The majority of current "
        "approaches are fundamentally reactive. They operate by observing the negative effects of learning (i.e., changes "
        "to important weights) and then applying a corrective measure. These methods, such as weight regularization, "
        "endeavor to preserve past knowledge by, for example, penalizing alterations to weights deemed critical for a "
        "preceding task. This often leads to over-conservatism, stifling the model's ability to learn genuinely novel "
        "concepts. A proactive paradigm that, instead of merely protecting the past, actively guides the learning process "
        "for optimal future performance remains a significantly underexplored domain. Such a paradigm would shift the "
        "focus from \"what to protect\" to \"how to learn\"."
    )
    
    document.add_paragraph('C. A New Proposition: From Bio-Inspiration to Cognitive Integration', style='Heading2IEEE')
    document.add_paragraph(
        "This paper asserts that while a substantial portion of AI has drawn inspiration from low-level biology "
        "(e.g., artificial neurons mimicking biological neurons), vast, untapped potential resides in the integration "
        "of higher-level principles from cognitive and developmental psychology [21, 22, 23, 24, 25]. Current machine "
        "learning practices often resemble a brain without a teacher or a student without a structured curriculum, "
        "relying on brute-force optimization over immense datasets. Human learning, by contrast, is frequently "
        "structured, supported, and guided by interaction."
    )
    document.add_paragraph(
        "The logical foundation of our proposition is articulated as follows:\n"
        "CL research is heavily concentrated on mechanistic solutions: regularization-based methods [30, 31], "
        "memory replay [32], and dynamic architectures. These represent bottom-up solutions that address the "
        "symptomatic manifestations of forgetting rather than its underlying cause—an inefficient learning process.\n"
        "A discrete body of research in educational psychology, pioneered by Vygotsky and Bruner, has formalized "
        "the mechanisms of effective human learning [21, 22, 23]. Key concepts in this domain include the "
        "Zone of Proximal Development (ZPD)—the conceptual space between what a learner can accomplish "
        "independently and what they can achieve with assistance—and Scaffolding—the temporary, tailored support "
        "provided by a \"More Knowledgeable Other\" (MKO) to facilitate the learner's traversal of their ZPD "
        "[17, 18, 19, 20]. This support is contingent on the learner's needs and is gradually faded as competence grows.\n"
        "Recent work has commenced the translation of scaffolding into computational terms for human-computer "
        "interaction, for instance, by employing reinforcement learning (RL) to develop adaptive tutoring systems "
        "that dynamically adjust problem difficulty or provide hints [19].\n"
        "The Conceptual Proposition: This work investigates the application of this concept not to a human learner, "
        "but to an AI model itself. It is possible to engineer a system wherein one AI component (a meta-controller) "
        "functions as the MKO, furnishing computational scaffolds to another AI component (the learner) when it "
        "enters its computational ZPD. This reframes CL from a problem of mitigating forgetting to one of "
        "optimizing guided knowledge acquisition. This constitutes a paradigm shift from a data-centric view to a "
        "process-centric view of lifelong learning."
    )
    
    document.add_paragraph('D. Contributions', style='Heading2IEEE')
    p = document.add_paragraph(style='List Bullet')
    p.add_run("The SMECI Framework:").bold = True
    p.add_run(" A novel, dual-network architecture that computationally implements the cognitive theory of scaffolded learning for an AI agent. This framework explicitly separates the \"learner\" from the \"teacher,\" enabling a new modality of meta-learning.")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("A Learnable ZPD Metric:").bold = True
    p.add_run(" The formulation of a novel, multifaceted metric that enables an AI to assess its own learning state in real-time. This metric, which combines indicators of uncertainty, gradient dynamics, and loss landscape geometry, allows the system to possess a form of self-awareness regarding its learning progress and difficulties.")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("A Meta-Learned Scaffolding Policy:").bold = True
    p.add_run(" The application of deep reinforcement learning to train a Scaffolding Network to dynamically select and apply the most efficacious computational support from a diverse action space. This moves beyond hard-coded rules, allowing the system to learn sophisticated, non-obvious teaching strategies optimized for long-term knowledge retention and acquisition efficiency.")
    p = document.add_paragraph(style='List Bullet')
    p.add_run("Guided Structural Plasticity:").bold = True
    p.add_run(" A novel algorithm for the consolidation of temporary scaffolds into the learner's permanent architecture. This ensures that network growth is not merely additive but is functionally meaningful, driven by demonstrated need and effectiveness, leading to more efficient and interpretable model expansion.")


    # --- Section II: Foundational Concepts and Related Work ---
    document.add_paragraph('II. FOUNDATIONAL CONCEPTS AND RELATED WORK', style='Heading1IEEE')
    document.add_paragraph('A. Paradigms in Continual Learning: A Critical Review', style='Heading2IEEE')
    document.add_paragraph(
        "This section presents a comprehensive review of the relevant literature, categorizing and critically analyzing existing CL methods."
    )
    p = document.add_paragraph()
    p.add_run("Regularization-Based Methods:").italic = True
    p.add_run(" These methods introduce a penalty term to the loss function to inhibit significant changes to parameters deemed important for previous tasks. A canonical example is Elastic Weight Consolidation (EWC) [29, 30, 31], which utilizes the Fisher Information Matrix to estimate parameter importance. The EWC loss is formulated as:")
    add_equation(document.add_paragraph(), "L(θ)=LB(θ) + ∑i (λ/2) Fi(θi − θA,i*)²", "1")
    document.add_paragraph(
        "where LB(θ) is the loss for the current task B, λ governs the importance of the old task, θA,i* are the optimal parameters from the preceding task A, and Fi represents the diagonal of the Fisher Information matrix, which approximates the precision of the posterior distribution [29]. Another prominent method is Learning without Forgetting (LwF), which uses knowledge distillation to ensure the outputs of the model for old task data remain similar after learning a new task. The principal limitation of such methods is their propensity for over-conservatism, which can impede plasticity and forward transfer, often leading to suboptimal solutions for new tasks."
    )
    
    # --- Section IV Table 1 ---
    document.add_paragraph('IV. CORE MECHANISMS AND MATHEMATICAL FORMULATIONS', style='Heading1IEEE') # Assuming this is the correct section for the table
    table1_headers = ["Action ai", "Cognitive Analogue [17, 18]", "Computational Mechanism", "Mathematical Impact on LN Update"]
    table1_data = [
        ("a0: No-Op", "Fading Support", "Continue standard training with no intervention.", "θt+1 = θt − η∇Lt"),
        ("a1: Attentional Focus", "Highlighting Key Features", "Instantiate and attach a temporary, lightweight self-attention module to a specific intermediate layer of the LN.", "Modifies the forward pass: hl′ = TempAttn(hl). Gradient flows through this temporary path, encouraging focus on relevant features."),
        ("a2: Task Decomposition", "Breaking Down the Task", "Create an auxiliary \"micro-network\" (e.g., a 2-layer MLP) that processes a problematic intermediate representation from the LN and is trained on a simplified auxiliary loss (e.g., predicting a subset of classes).", "Adds an auxiliary loss term: Ltotal = Lmain + λauxLaux, creating an additional gradient signal to guide feature learning."),
        ("a3: Plasticity Regulation", "Guided Practice", "Temporarily modulate the learning rates on a per-layer basis. For layers identified as critical for past tasks (using Fisher Information), the learning rate is decreased, while for newly added or less critical layers, it is increased.", "Introduces a per-layer or per-parameter learning rate matrix ηt, such that Δθt = −ηt ⊙ ∇Lt."),
        ("a4: Structural Expansion", "Providing New Tools", "Trigger the DEN mechanism [34, 35, 36, 37, 38, 39] to add a small number of new neurons to a layer exhibiting high semantic drift or capacity saturation, as identified by the SN.", "Increases the dimensionality of the weight matrices Wl, providing new capacity for learning."),
        ("a5: Prioritized Rehearsal", "Reminding of Past Concepts", "If a rehearsal buffer is used, this action temporarily increases the sampling rate of exemplars from a past task deemed most similar to the current one, to explicitly manage inter-task interference.", "Modifies the data batch Dt to include a higher proportion of specific past data."),
    ]
    add_ieee_table(document, "TABLE I\nTHE SCAFFOLDING ACTION SPACE", table1_headers, table1_data)

    # --- Section VII: Conclusion ---
    document.add_paragraph('VII. CONCLUSION AND FUTURE WORK', style='Heading1IEEE')
    document.add_paragraph(
        "This paper has introduced the SMECI framework, a novel, cognitively-grounded solution to the problem of lifelong learning. "
        "By computationally operationalizing the concepts of the Zone of Proximal Development and cognitive scaffolding, SMECI proposes a "
        "paradigm shift from reactive knowledge preservation to proactive, guided learning. The key contributions—a dual-network "
        "architecture, a self-assessment metric for learning, a meta-learned instructional policy, and guided structural plasticity—were detailed. "
        "The promising, albeit preliminary, implications of this approach suggest a path toward more adaptive, efficient, and robust AI systems."
    )
    
    # --- Appendix & References ---
    document.add_paragraph('APPENDIX', style='Heading1IEEE')
    document.add_paragraph('A. Meta-Training Algorithm for the Scaffolding Network', style='Heading2IEEE')
    document.add_paragraph(
        "# Algorithm 1: SMECI Meta-Training Loop\n"
        "# Initialize Learner Network (LN) parameters: theta_ln\n"
        "# ... (rest of the algorithm text) ..."
    )
    
    document.add_paragraph('REFERENCES', style='Heading1IEEE')
    p = document.add_paragraph("[1] G. O. Young, “Synthetic structure of industrial plastics,” in Plastics, 2nd ed., vol. 3, J. Peters, Ed. New York, NY, USA: McGraw-Hill, 1964, pp. 15–64.", style='Reference')
    p = document.add_paragraph("[2] W.-K. Chen, Linear Networks and Systems. Belmont, CA, USA: Wadsworth, 1993, pp. 123–135.", style='Reference')
    # ... Add all other references here ...
    p = document.add_paragraph("[15] IEEE, IEEE Editorial Style Manual for Authors.", style='Reference')


    # Save the document
    file_path = "SMECI_IEEE_Formatted_Article.docx"
    document.save(file_path)
    print(f"مقاله با موفقیت در فایل '{file_path}' ذخیره شد.")

if __name__ == '__main__':
    create_full_document()

